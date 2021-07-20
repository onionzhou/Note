from docx import Document
from docx.shared import Pt # 设置字号

document = Document()

# 表格1
# 标题1，字体为等线，且为斜体
title1 = document.add_heading(u'1.表格1', level=1)
title1.style.font.name = u'等线' # 设置中文字体前面要有u
title1.italic = True

table1 = document.add_table(rows=3,cols=7) # 3行7列
# 设置表格标题栏
for i in range(7):
    table1.cell(0,i).text = str(i+1)

table1.style="Light Shading" # 风格为三线表


# 表格2
# 标题2，字号为12，且为粗体
title2 = document.add_paragraph(u'2.表格2',style="heading 1")
title2.style.font.size = Pt(12)
title2.bold = True

table2 = document.add_table(rows=3,cols=7) # 3行7列
headLine = ["A","B","C","D","E","F","G"]
# 设置表格标题栏
for i in range(7):
    table2.cell(0,i).text = headLine[i]

table2.style="Light Shading" # 风格为三线表

# 储存
document.save("test.docx")