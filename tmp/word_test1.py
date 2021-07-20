#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
@author: zhouey
@file: word_test1.py
@time: 2020/10/26 14:33
"""
from docx import Document
from docx.shared import Inches


def main():
    document = Document()

    document.add_heading('渤海api文档', 0)
    document.add_heading('1部门', level=1)
    document.add_heading('新建部门', level=2)
    document.add_heading('输入', level=3)
    document.add_paragraph(
        'url: https://xxxx', style='List Bullet'
    )
    document.add_paragraph(
        '方法：post', style='List Bullet'
    )
    document.add_paragraph(
        '参数:', style='List Bullet'
    )
    document.add_heading('输出', level=3)

    document.save('demo1.docx')


if __name__ == '__main__':
    main()
